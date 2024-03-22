from docx import Document

def create_document(doc_path, doc_name, num_pages):
    document = Document()

    for _ in range(num_pages):
        # Add a page break for all pages except the first one
        if document.paragraphs:
            document.add_page_break()

        # Add a table with 5 rows and 2 columns to each page
        table = document.add_table(rows=5, cols=2)

        # Add some dummy data to the table
        for row in table.rows:
            for cell in row.cells:
                cell.text = "Sample Data"

    # Save the document with the given name and path
    document.save(f"{doc_path}/{doc_name}.docx")

if __name__ == "__main__":
    # Get user input
    doc_path = input("Enter document path (e.g., C:/Users/37251/Desktop/): ")
    doc_name = input("Enter document name: ")
    num_pages = int(input("Enter number of pages: "))

    # Create the document
    create_document(doc_path, doc_name, num_pages)

    print(f"Document '{doc_name}.docx' created successfully at {doc_path}.")
