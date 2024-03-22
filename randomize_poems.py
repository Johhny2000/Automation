import re
import random
from docx import Document


def randomize_number(match, used_numbers):
    number = int(match.group(1))
    while True:
        randomized_number = random.randint(1, 40)
        if randomized_number not in used_numbers:
            used_numbers.add(randomized_number)
            return f'({randomized_number})'


def randomize_poem_numbers(poem_count):
    poem_numbers = list(range(1, poem_count + 1))
    random.shuffle(poem_numbers)
    return poem_numbers

def scan_and_randomize(doc_path):
    document = Document(doc_path)

    # Ask the user for the number of poems to randomize
    poem_count = int(input("Enter the number of poems to randomize: "))
    selected_poem_numbers = randomize_poem_numbers(poem_count)

    poem_index = 0

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if re.search(r'\(\d+\)', text):
                        updated_text = re.sub(r'\((\d+)\)', lambda match: f'({selected_poem_numbers[poem_index]})', text)
                        poem_index += 1
                        paragraph.clear()
                        paragraph.add_run(updated_text)

    updated_doc_path = doc_path.replace('.docx', '_randomized.docx')
    document.save(updated_doc_path)

    print(f"Randomized document saved at: {updated_doc_path}")


if __name__ == "__main__":
    doc_path = "C:/Users/37251/Desktop/Doc/luuletused_2.docx"
    scan_and_randomize(doc_path)
